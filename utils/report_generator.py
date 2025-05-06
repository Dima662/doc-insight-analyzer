from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_COLOR_INDEX

def generate_report_docx(summary, entities, original_text):
    doc = Document()
    doc.add_heading("Doc Insight Analyzer Report", 0)

    doc.add_heading("Summary", level=1)
    doc.add_paragraph(summary)

    doc.add_heading("Named Entities", level=1)
    for label, items in entities.items():
        doc.add_paragraph(f"{label}: {', '.join(items)}", style="List Bullet")

    doc.add_heading("Highlighted Text", level=1)
    para = doc.add_paragraph()

    words = original_text.split()
    for word in words:
        stripped = word.strip('.,?!()[]"\'')
        matched = any(stripped == ent for group in entities.values() for ent in group)
        run = para.add_run(word + " ")
        if matched:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        run.font.size = Pt(10)

    return doc
